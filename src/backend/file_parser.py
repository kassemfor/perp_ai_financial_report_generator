"""
Multi-format file parser: CSV, TXT, PDF, DOCX, XLSX.
"""

import re
import shutil
import pandas as pd
from pathlib import Path
from statistics import mean
from typing import Any, Dict, List, Optional, Tuple

class FileParser:
    """Parse multiple file formats and extract structured data."""

    def __init__(self):
        self.image_formats = ["png", "jpg", "jpeg", "tiff", "bmp", "webp"]
        self.supported_formats = ["csv", "txt", "pdf", "docx", "xlsx", *self.image_formats]

    def parse(self, file_path: str, enable_ocr: bool = True, ocr_language: str = "eng") -> Dict[str, Any]:
        """Parse file and return structured data."""
        file_path = Path(file_path)
        ext = file_path.suffix.lower().lstrip(".")

        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {ext}")

        if ext == "csv":
            return self._parse_csv(file_path)
        if ext == "txt":
            return self._parse_txt(file_path)
        if ext == "pdf":
            return self._parse_pdf(file_path, enable_ocr=enable_ocr, ocr_language=ocr_language)
        if ext == "docx":
            return self._parse_docx(file_path)
        if ext == "xlsx":
            return self._parse_xlsx(file_path)
        return self._parse_image(file_path, ocr_language=ocr_language)

    def _parse_csv(self, file_path: Path) -> Dict[str, Any]:
        """Parse CSV file."""
        try:
            df = pd.read_csv(file_path)
            records = df.to_dict("records")
            flattened_text = self._records_to_text(records)
            confidence = self._estimate_extraction_confidence(text=flattened_text, ocr_used=False)
            return {
                "format": "csv",
                "file_name": file_path.name,
                "rows": len(df),
                "columns": list(df.columns),
                "data": records,
                "summary": df.describe(include="all").to_dict(),
                "dtypes": df.dtypes.astype(str).to_dict(),
                "text": flattened_text,
                "ocr": {"enabled": False, "used": False, "status": "not_required", "language": "eng"},
                "extraction_source": "structured_csv",
                "extraction_confidence": confidence,
                "rfs_statement": self._build_rfs_statement(
                    text=flattened_text,
                    source="structured_csv",
                    confidence=confidence,
                ),
            }
        except Exception as e:
            raise RuntimeError(f"CSV parsing failed: {str(e)}")

    def _parse_xlsx(self, file_path: Path) -> Dict[str, Any]:
        """Parse Excel file."""
        try:
            xls = pd.ExcelFile(file_path)
            sheets: Dict[str, Dict[str, Any]] = {}
            aggregate_records: List[Dict[str, Any]] = []
            for sheet in xls.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet)
                records = df.to_dict("records")
                sheets[sheet] = {
                    "rows": len(df),
                    "columns": list(df.columns),
                    "data": records,
                    "summary": df.describe(include="all").to_dict(),
                }
                # Keep bounded samples for RFS normalization context.
                for row in records[:80]:
                    normalized_row = {"sheet": sheet}
                    normalized_row.update(row)
                    aggregate_records.append(normalized_row)

            flattened_text = self._records_to_text(aggregate_records, max_rows=400)
            confidence = self._estimate_extraction_confidence(text=flattened_text, ocr_used=False)
            return {
                "format": "xlsx",
                "file_name": file_path.name,
                "sheets": sheets,
                "sheet_names": list(xls.sheet_names),
                "text": flattened_text,
                "ocr": {"enabled": False, "used": False, "status": "not_required", "language": "eng"},
                "extraction_source": "structured_xlsx",
                "extraction_confidence": confidence,
                "rfs_statement": self._build_rfs_statement(
                    text=flattened_text,
                    source="structured_xlsx",
                    confidence=confidence,
                ),
            }
        except Exception as e:
            raise RuntimeError(f"XLSX parsing failed: {str(e)}")

    def _parse_txt(self, file_path: Path) -> Dict[str, Any]:
        """Parse text file."""
        try:
            content = file_path.read_text()
            lines = content.split("\n")
            # Try to detect table structure
            tables = self._extract_tables_from_text(content)
            return {
                "format": "txt",
                "file_name": file_path.name,
                "content": content,
                "line_count": len(lines),
                "char_count": len(content),
                "tables": tables,
                "rfs_statement": self._build_rfs_statement(
                    text=content,
                    source="plain_text",
                    confidence=self._estimate_extraction_confidence(content, ocr_used=False),
                ),
            }
        except Exception as e:
            raise RuntimeError(f"TXT parsing failed: {str(e)}")

    def _parse_pdf(self, file_path: Path, enable_ocr: bool = True, ocr_language: str = "eng") -> Dict[str, Any]:
        """Parse PDF file."""
        try:
            import pypdf
        except ImportError:
            raise ImportError("pypdf required: pip install pypdf")

        try:
            pdf = pypdf.PdfReader(file_path)
            page_count = len(pdf.pages)
            text_chunks: List[str] = []
            metadata = pdf.metadata

            for page in pdf.pages:
                extracted = (page.extract_text() or "").strip()
                if extracted:
                    text_chunks.append(extracted)

            text = "\n\n".join(text_chunks)
            extraction_source = "embedded_text"
            ocr_used = False
            ocr_info: Dict[str, Any] = {"status": "not_required", "language": ocr_language}
            ocr_confidence = 0.0

            if enable_ocr and self._should_ocr_pdf(text, page_count):
                ocr_text, ocr_details = self._ocr_pdf(file_path, ocr_language=ocr_language)
                if ocr_text.strip():
                    text = ocr_text
                    extraction_source = "ocr_pdf"
                    ocr_used = True
                    ocr_confidence = float(ocr_details.get("confidence", 0.0))
                    ocr_info = {
                        "status": "success",
                        "language": ocr_language,
                        "backend": ocr_details.get("backend", "unknown"),
                        "confidence": ocr_confidence,
                    }
                else:
                    ocr_info = {
                        "status": ocr_details.get("status", "unavailable"),
                        "language": ocr_language,
                        "error": ocr_details.get("error", "OCR engine unavailable"),
                    }

            tables = self._extract_tables_from_text(text)
            confidence = self._estimate_extraction_confidence(
                text=text,
                ocr_used=ocr_used,
                ocr_confidence=ocr_confidence if ocr_used else None,
            )
            safe_metadata = {str(k): str(v) for k, v in dict(metadata or {}).items()}

            return {
                "format": "pdf",
                "file_name": file_path.name,
                "pages": page_count,
                "text": text,
                "metadata": safe_metadata,
                "tables": tables,
                "ocr": {
                    "enabled": enable_ocr,
                    "used": ocr_used,
                    **ocr_info,
                },
                "extraction_source": extraction_source,
                "extraction_confidence": confidence,
                "rfs_statement": self._build_rfs_statement(
                    text=text,
                    source=extraction_source,
                    confidence=confidence,
                ),
            }
        except Exception as e:
            raise RuntimeError(f"PDF parsing failed: {str(e)}")

    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        try:
            doc = Document(file_path)
            text = ""
            tables_data = []

            for para in doc.paragraphs:
                text += para.text + "\n"

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables_data.append(table_data)

            confidence = self._estimate_extraction_confidence(text=text, ocr_used=False)
            return {
                "format": "docx",
                "file_name": file_path.name,
                "text": text,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(tables_data),
                "tables": tables_data,
                "extraction_source": "document_text",
                "extraction_confidence": confidence,
                "rfs_statement": self._build_rfs_statement(
                    text=text,
                    source="document_text",
                    confidence=confidence,
                ),
            }
        except Exception as e:
            raise RuntimeError(f"DOCX parsing failed: {str(e)}")

    def _parse_image(self, file_path: Path, ocr_language: str = "eng") -> Dict[str, Any]:
        """Parse image files using OCR."""
        try:
            from PIL import Image
        except ImportError:
            raise ImportError("Pillow required for image parsing: pip install pillow")

        try:
            with Image.open(file_path) as image:
                width, height = image.size
                mode = image.mode
                text, confidence, error = self._ocr_image_object(image, ocr_language=ocr_language)

            tables = self._extract_tables_from_text(text)
            extraction_source = "ocr_image" if text.strip() else "image_unreadable"
            extraction_confidence = self._estimate_extraction_confidence(
                text=text,
                ocr_used=True,
                ocr_confidence=confidence,
            )

            return {
                "format": file_path.suffix.lower().lstrip("."),
                "file_name": file_path.name,
                "width": width,
                "height": height,
                "mode": mode,
                "text": text,
                "tables": tables,
                "ocr": {
                    "enabled": True,
                    "used": True,
                    "status": "success" if text.strip() else "failed",
                    "language": ocr_language,
                    "confidence": confidence,
                    "error": error,
                },
                "extraction_source": extraction_source,
                "extraction_confidence": extraction_confidence,
                "rfs_statement": self._build_rfs_statement(
                    text=text,
                    source=extraction_source,
                    confidence=extraction_confidence,
                ),
            }
        except Exception as e:
            raise RuntimeError(f"Image parsing failed: {str(e)}")

    def _extract_tables_from_text(self, text: str) -> List[List[List[str]]]:
        """Attempt to extract table-like structures from text."""
        if not text:
            return []

        lines = text.split("\n")
        tables: List[List[List[str]]] = []
        current_table: List[List[str]] = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if len(current_table) >= 2:
                    tables.append(current_table)
                current_table = []
                continue

            delimiter = None
            if "|" in line:
                delimiter = "|"
            elif "\t" in line:
                delimiter = "\t"
            elif line.count(",") >= 2:
                delimiter = ","

            if delimiter:
                row = [cell.strip() for cell in line.split(delimiter)]
                if len([cell for cell in row if cell]) >= 2:
                    current_table.append(row)
                    continue

            if len(current_table) >= 2:
                tables.append(current_table)
            current_table = []

        if len(current_table) >= 2:
            tables.append(current_table)

        return tables

    def _records_to_text(self, records: List[Dict[str, Any]], max_rows: int = 250) -> str:
        """Flatten structured rows into text lines for unified RFS normalization."""
        lines: List[str] = []
        for row in records[:max_rows]:
            if not isinstance(row, dict):
                continue
            parts = []
            for key, value in row.items():
                if value is None:
                    continue
                value_text = str(value).strip()
                if not value_text:
                    continue
                parts.append(f"{key}: {value_text}")
            if parts:
                lines.append(" | ".join(parts))
        return "\n".join(lines)

    def _should_ocr_pdf(self, extracted_text: str, page_count: int) -> bool:
        """Determine whether PDF likely needs OCR."""
        compact_length = len(re.sub(r"\s+", "", extracted_text or ""))
        min_expected = max(120, page_count * 60)
        return compact_length < min_expected

    def _ocr_pdf(self, file_path: Path, ocr_language: str = "eng") -> Tuple[str, Dict[str, Any]]:
        """OCR a PDF by rendering pages to images and running OCR."""
        images, backend, render_error = self._render_pdf_to_images(file_path)
        if not images:
            return "", {
                "status": "unavailable",
                "error": render_error or "No PDF render backend available",
                "backend": backend or "none",
            }

        text_parts: List[str] = []
        confidences: List[float] = []
        ocr_errors: List[str] = []

        for image in images:
            text, confidence, error = self._ocr_image_object(image, ocr_language=ocr_language)
            if text.strip():
                text_parts.append(text.strip())
            if confidence > 0:
                confidences.append(confidence)
            if error:
                ocr_errors.append(error)

        if not text_parts:
            return "", {
                "status": "failed",
                "error": "; ".join(sorted(set(ocr_errors))) or "No text detected in OCR",
                "backend": backend or "unknown",
            }

        return "\n\n".join(text_parts), {
            "status": "success",
            "backend": backend or "unknown",
            "confidence": round(mean(confidences), 3) if confidences else 0.0,
            "error": "; ".join(sorted(set(ocr_errors))) if ocr_errors else "",
        }

    def _render_pdf_to_images(self, file_path: Path) -> Tuple[List[Any], str, str]:
        """Render PDF pages to images using available backend."""
        try:
            import pypdfium2 as pdfium

            images: List[Any] = []
            pdf_doc = pdfium.PdfDocument(str(file_path))
            for index in range(len(pdf_doc)):
                page = pdf_doc[index]
                bitmap = page.render(scale=2.0)
                pil_image = bitmap.to_pil()
                images.append(pil_image)
                page.close()
            pdf_doc.close()
            return images, "pypdfium2", ""
        except Exception as pypdfium_error:
            pypdfium_message = f"pypdfium2 unavailable: {str(pypdfium_error)}"

        try:
            from pdf2image import convert_from_path

            images = convert_from_path(str(file_path), dpi=250)
            return images, "pdf2image", ""
        except Exception as pdf2image_error:
            return [], "", f"{pypdfium_message}; pdf2image unavailable: {str(pdf2image_error)}"

    def _ocr_image_object(self, image: Any, ocr_language: str = "eng") -> Tuple[str, float, str]:
        """Run OCR for a loaded image object."""
        try:
            import pytesseract
        except ImportError:
            return "", 0.0, "pytesseract package not installed"

        if shutil.which("tesseract") is None:
            return "", 0.0, "tesseract binary not installed"

        try:
            text = pytesseract.image_to_string(image, lang=ocr_language)
            details = pytesseract.image_to_data(
                image,
                lang=ocr_language,
                output_type=pytesseract.Output.DICT,
            )
            conf_scores = []
            for value in details.get("conf", []):
                try:
                    score = float(value)
                    if score >= 0:
                        conf_scores.append(score / 100.0)
                except (ValueError, TypeError):
                    continue
            confidence = round(mean(conf_scores), 3) if conf_scores else 0.0
            return text, confidence, ""
        except Exception as e:
            return "", 0.0, str(e)

    def _estimate_extraction_confidence(
        self,
        text: str,
        ocr_used: bool,
        ocr_confidence: Optional[float] = None,
    ) -> float:
        """Estimate confidence for parsed text extraction."""
        if not text.strip():
            return 0.0

        compact_length = len(re.sub(r"\s+", "", text))
        base = min(0.97, 0.45 + (compact_length / 6000.0))
        if ocr_used:
            if ocr_confidence is not None and ocr_confidence > 0:
                return round(min(0.95, max(0.35, ocr_confidence)), 3)
            return round(min(0.9, base * 0.85), 3)
        return round(base, 3)

    def _build_rfs_statement(self, text: str, source: str, confidence: float) -> Dict[str, Any]:
        """Build normalized statement output aligned to an RFS-style schema."""
        normalized_text = text or ""
        numeric_pattern = re.compile(r"[$€£]?\s*-?\(?\d[\d,]*(?:\.\d+)?%?\)?")
        period_pattern = re.compile(
            r"(?:Q[1-4]\s*20\d{2}|(?:19|20)\d{2}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+(?:19|20)\d{2})",
            re.IGNORECASE,
        )

        statement_lines: List[Dict[str, Any]] = []
        for raw_line in normalized_text.splitlines():
            line = re.sub(r"\s+", " ", raw_line).strip()
            if len(line) < 4:
                continue
            amount_match = numeric_pattern.search(line)
            if not amount_match:
                continue
            line_item = line[: amount_match.start()].strip(":- ")
            if len(line_item) < 2:
                line_item = "Unlabeled Line Item"
            period_match = period_pattern.search(line)
            rfs_code = self._classify_rfs_code(line_item)
            statement_lines.append(
                {
                    "line_item": line_item,
                    "value": amount_match.group().replace(" ", ""),
                    "period": period_match.group(0) if period_match else "",
                    "source": source,
                    "confidence": round(confidence, 3),
                    "rfs_code": rfs_code,
                }
            )
            if len(statement_lines) >= 40:
                break

        recognized_codes = sorted({line["rfs_code"] for line in statement_lines if line.get("rfs_code") != "OTHER"})
        statement_type = self._infer_statement_type(recognized_codes)
        checks = {
            "minimum_text_length": len(normalized_text.strip()) >= 80,
            "contains_numeric_values": bool(numeric_pattern.search(normalized_text)),
            "contains_statement_lines": len(statement_lines) > 0,
            "traceable_source": source in {"embedded_text", "plain_text", "document_text", "ocr_pdf", "ocr_image"},
            "contains_recognized_rfs_codes": len(recognized_codes) > 0,
        }
        quality_score = round((sum(checks.values()) / len(checks)) * 100, 1)
        warnings = []
        if not checks["minimum_text_length"]:
            warnings.append("Low extracted text volume")
        if not checks["contains_statement_lines"]:
            warnings.append("No structured statement lines detected")
        if not checks["contains_recognized_rfs_codes"]:
            warnings.append("No recognized RFS financial categories detected")

        return {
            "standard": "RFS-STATEMENT-v1",
            "schema_version": "1.1",
            "status": "pass" if quality_score >= 75 else "needs_review",
            "quality_score": quality_score,
            "confidence": round(confidence, 3),
            "checks": checks,
            "warnings": warnings,
            "document_profile": {
                "statement_type": statement_type,
                "recognized_rfs_codes": recognized_codes,
            },
            "summary": {
                "line_items_detected": len(statement_lines),
                "characters": len(normalized_text),
                "source": source,
            },
            "statement_lines": statement_lines,
        }

    def _classify_rfs_code(self, line_item: str) -> str:
        """Map line item labels to canonical RFS-like financial codes."""
        item = (line_item or "").lower()
        patterns = [
            (r"revenue|sales|turnover", "REV"),
            (r"cost of goods|cogs|cost of sales", "COGS"),
            (r"gross profit|gross margin", "GROSS_PROFIT"),
            (r"operating income|operating profit|ebit", "OPERATING_INCOME"),
            (r"net income|net profit|profit after tax|earnings", "NET_INCOME"),
            (r"expense|opex|operating expense", "EXPENSE"),
            (r"asset", "ASSETS"),
            (r"liabilit", "LIABILITIES"),
            (r"equity|shareholder", "EQUITY"),
            (r"cash flow|cashflow", "CASH_FLOW"),
            (r"ebitda", "EBITDA"),
            (r"tax", "TAX"),
        ]
        for pattern, code in patterns:
            if re.search(pattern, item):
                return code
        return "OTHER"

    def _infer_statement_type(self, recognized_codes: List[str]) -> str:
        """Infer likely statement type from recognized line item codes."""
        code_set = set(recognized_codes)
        if code_set & {"REV", "COGS", "GROSS_PROFIT", "OPERATING_INCOME", "NET_INCOME"}:
            return "income_statement"
        if code_set & {"ASSETS", "LIABILITIES", "EQUITY"}:
            return "balance_sheet"
        if code_set & {"CASH_FLOW"}:
            return "cash_flow_statement"
        return "general_statement"
